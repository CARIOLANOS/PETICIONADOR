import streamlit as st
import fitz  # PyMuPDF
import spacy
import mimetypes
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Carrega modelo de linguagem natural para português
nlp = spacy.load("pt_core_news_sm")

def apply_style(paragraph, style):
    run = paragraph.runs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run.font.name = 'Times New Roman'
    style = style.lower()

    if style == 'corpo':
        run.font.size = Pt(12)
        paragraph.paragraph_format.line_spacing = Pt(18)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.first_line_indent = Cm(1.25)

    elif style == 'citação':
        run.font.size = Pt(10)
        paragraph.paragraph_format.left_indent = Cm(4)
        paragraph.paragraph_format.line_spacing = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    elif style == 'título':
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_after = Pt(6)

    elif style in ['endereçamento', 'pedidos']:
        run.font.size = Pt(12)
        paragraph.paragraph_format.line_spacing = Pt(18)
        paragraph.paragraph_format.space_after = Pt(6)
        if style == 'pedidos':
            paragraph.paragraph_format.first_line_indent = Cm(1.25)

def read_txt(file):
    return file.read().decode("utf-8")

def read_docx(file):
    doc = Document(file)
    return '\n'.join(p.text for p in doc.paragraphs)

def read_pdf(file):
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in pdf:
        text += page.get_text()
    return text

def analyze_style(text):
    doc = nlp(text)
    if text.isupper():
        return "Título"
    elif any(t.text in ['"', '“', '”'] for t in doc):
        return "Citação"
    elif len(text.strip().split()) <= 5:
        return "Endereçamento"
    elif text.strip().endswith(":"):
        return "Pedidos"
    else:
        return "Corpo"

def create_docx(paragraphs, filename):
    doc = Document()
    doc.add_heading('Petição', 0)

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    for para, style in paragraphs:
        p = doc.add_paragraph(para)
        apply_style(p, style)

    doc.save(filename)

def main():
    st.set_page_config(page_title="Formatador ABNT", layout="wide")
    st.title("📄 Formatador de Peças Processuais - ABNT")
    st.header("by Cariolano")

    uploaded_file = st.file_uploader("📎 Selecione seu arquivo (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])

    if uploaded_file:
        st.info("🔍 Lendo arquivo e processando...")

        file_type, _ = mimetypes.guess_type(uploaded_file.name)
        if file_type == "text/plain":
            text = read_txt(uploaded_file)
        elif file_type == "application/pdf":
            text = read_pdf(uploaded_file)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = read_docx(uploaded_file)
        else:
            st.error("Tipo de arquivo não suportado.")
            return

        paragraphs_raw = [p for p in text.split("\n") if p.strip()]
        formatted_paragraphs = []

        st.subheader("🧠 Sugestão automática de estilo ABNT")
        for i, para in enumerate(paragraphs_raw):
            style_suggested = analyze_style(para)
            st.markdown(f"**Parágrafo {i+1}:**")
            st.text_area("Texto:", value=para, height=100, key=f"text_{i}", disabled=True)
            style_final = st.selectbox(
                "Estilo sugerido (você pode ajustar):",
                ["Corpo", "Citação", "Título", "Endereçamento", "Pedidos"],
                index=["Corpo", "Citação", "Título", "Endereçamento", "Pedidos"].index(style_suggested),
                key=f"style_{i}"
            )
            formatted_paragraphs.append((para, style_final))

        if st.button("👁️ Pré-visualizar Documento"):
            st.subheader("📄 Pré-visualização")
            for i, (para, style) in enumerate(formatted_paragraphs):
                st.markdown(f"**{style}**")
                st.write(para)
                st.markdown("---")

        if st.button("💾 Gerar e Baixar Documento"):
            output_filename = "peticao_formatada.docx"
            create_docx(formatted_paragraphs, output_filename)
            with open(output_filename, "rb") as file:
                st.download_button(
                    label="📥 Baixar Documento",
                    data=file,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
